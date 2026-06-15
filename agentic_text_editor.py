
from typing import Type
from crewai import Agent,Task,Crew,Process,LLM
from docx import Document 
import nltk
from nltk import word_tokenize,sent_tokenize
import streamlit as st

API_KEY=st.secrets.api_key
API_URL = 'https://api.openai.com/v1/chat/completions'
OPENAI_API_KEY="sk-proj-your key"
llm=LLM(model='openai/o1-mini',temperature=0.75,max_tokens=16000,api_key=OPENAI_API_KEY)



def read_docx(file_path):
    """ Extract text from DOCX file. """
    req_w,l=0,0
    doc = Document(file_path)
    for para in doc.paragraphs:
        text+=para.text+' '
        words=nltk.word_tokenize(para.text)
        l+=len(words)         
    req_w = round(l*1.05*3)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# instructions=""
def load_instructions(filename):
    """ Load editing instructions from a file. """
    try:
        with open(filename, 'r') as file:
            data = file.read()
            return data
            st.write('file returned')
    except Exception as e:
        st.error('An error occcured:{e}')
        return "Default instructions if file doesn't exist."

def editor(llm,filename,options,report_features,edits):
    instructions = load_instructions(options)
    combined_text = instructions + " " + " ".join([report_features[feature] for feature in edits])

    res_agent=Agent(
        role = "Research Paper Editor",
        goal = """To edit the given research paper strictly as per the provided instructions.""",
        backstory = "You are an expert at editing research papers using the given instructions and have a keen eye for detail and keeping up with the provided guidelines for editing.",
        llm=llm,
        max_iterations=2,
        )

    ppt_task = Task(
        description = f"""Edit the research paper as per the provided instructions.
        The research paper: {read_docx(filename)}
        The instructions are: {combined_text}""",
        agent=res_agent,
        expected_output = "Edited research paper",
    )

    ppt_crew = Crew(agents=[res_agent], tasks=[ppt_task],process=Process.sequential)

    result=ppt_crew.kickoff()
